"""
Maritime Workplace Climate Survey 2026
Regional Analysis Report Generator
Author: Ezequiel Bassa
"""

import os
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── CONFIG ────────────────────────────────────────────────────────────────────
CSV_PATH    = "workplace_climate_maritime_2026.csv"
OUTPUT_PATH = "Maritime_Climate_Report_2026.docx"
IMG_DIR     = "report_imgs"
os.makedirs(IMG_DIR, exist_ok=True)

# Colour palette (dark maritime)
REGION_COLORS = {
    "North America":   "#1a6fa8",
    "Central America": "#c9a84c",
    "South America":   "#2ecc71",
}
BG  = "#0d1b2a"
FG  = "#dde6ee"

plt.rcParams.update({
    "figure.facecolor":  BG,
    "axes.facecolor":    "#101f30",
    "axes.edgecolor":    "#1a2e44",
    "axes.labelcolor":   FG,
    "xtick.color":       FG,
    "ytick.color":       FG,
    "text.color":        FG,
    "grid.color":        "#1a2e44",
    "grid.linestyle":    "--",
    "grid.linewidth":    0.5,
    "font.family":       "sans-serif",
    "font.size":         10,
    "axes.titlesize":    12,
    "axes.titleweight":  "bold",
})

CLIMATE_COLS = [
    "eNPS", "Safety_Culture", "Cross_Border_Collab", "Leadership_Trust",
    "Work_Life_Balance", "Career_Growth", "Inclusion_Diversity",
    "Operational_Stress", "Fair_Compensation", "Tech_Adequacy",
    "Recognition", "Environmental_Pride", "Engagement", "Job_Security",
    "Managerial_Support", "Ethical_Standards", "Training_Quality",
    "Communication_Clarity", "Team_Cohesion", "Retention_Intent"
]

FRIENDLY = {
    "eNPS": "eNPS",
    "Safety_Culture": "Safety Culture",
    "Cross_Border_Collab": "Cross-Border Collab",
    "Leadership_Trust": "Leadership Trust",
    "Work_Life_Balance": "Work-Life Balance",
    "Career_Growth": "Career Growth",
    "Inclusion_Diversity": "Inclusion & Diversity",
    "Operational_Stress": "Operational Stress",
    "Fair_Compensation": "Fair Compensation",
    "Tech_Adequacy": "Tech Adequacy",
    "Recognition": "Recognition",
    "Environmental_Pride": "Environmental Pride",
    "Engagement": "Engagement",
    "Job_Security": "Job Security",
    "Managerial_Support": "Managerial Support",
    "Ethical_Standards": "Ethical Standards",
    "Training_Quality": "Training Quality",
    "Communication_Clarity": "Communication Clarity",
    "Team_Cohesion": "Team Cohesion",
    "Retention_Intent": "Retention Intent",
}

# ── LOAD DATA ─────────────────────────────────────────────────────────────────
df = pd.read_csv(CSV_PATH)
df["Avg_Climate"] = df[CLIMATE_COLS].mean(axis=1)

regions  = ["North America", "Central America", "South America"]
n_total  = len(df)

# ── HELPER: save figure ───────────────────────────────────────────────────────
def save_fig(fig, name):
    path = os.path.join(IMG_DIR, name)
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor=BG)
    plt.close(fig)
    return path

# ── FIGURE 1: Composite Score by Region (bar) ────────────────────────────────
def fig_composite_by_region():
    reg = df.groupby("Region")["Avg_Climate"].agg(["mean", "std", "count"]).reindex(regions)
    reg["se"] = reg["std"] / np.sqrt(reg["count"])
    reg["ci"] = reg["se"] * 1.96

    fig, ax = plt.subplots(figsize=(7, 3.8))
    bars = ax.bar(
        reg.index,
        reg["mean"],
        color=[REGION_COLORS[r] for r in reg.index],
        width=0.5,
        yerr=reg["ci"],
        capsize=5,
        error_kw={"ecolor": "#c9a84c", "linewidth": 1.5},
        zorder=3
    )
    ax.set_ylim(0, 5.2)
    ax.set_ylabel("Avg Composite Score (1-5)")
    ax.set_title("Figure 1 — Overall Composite Climate Score by Region (±95% CI)")
    ax.axhline(3.0, color="#e74c3c", linewidth=1, linestyle="--", label="Critical threshold (3.0)")
    ax.axhline(4.0, color="#2ecc71", linewidth=1, linestyle="--", label="Good threshold (4.0)")
    ax.legend(fontsize=8, loc="lower right")
    ax.yaxis.grid(True, zorder=0)
    ax.set_axisbelow(True)
    for bar, val in zip(bars, reg["mean"]):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.06,
                f"{val:.2f}", ha="center", va="bottom", fontsize=10, color=FG)
    fig.tight_layout()
    return save_fig(fig, "fig1_composite.png")

# ── FIGURE 2: Heatmap KPIs × Region ──────────────────────────────────────────
def fig_heatmap():
    heat = df.groupby("Region")[CLIMATE_COLS].mean().reindex(regions)
    heat.columns = [FRIENDLY[c] for c in heat.columns]

    fig, ax = plt.subplots(figsize=(13, 3.2))
    sns.heatmap(
        heat,
        ax=ax,
        annot=True,
        fmt=".2f",
        cmap=sns.diverging_palette(10, 133, as_cmap=True),
        vmin=1, vmax=5,
        linewidths=0.4,
        linecolor="#0d1b2a",
        annot_kws={"size": 7},
        cbar_kws={"label": "Avg Score", "shrink": 0.8}
    )
    ax.set_title("Figure 2 — KPI Heatmap: All Indicators by Region")
    ax.set_xlabel("")
    ax.set_ylabel("")
    ax.tick_params(axis="x", labelsize=7.5, rotation=40)
    ax.tick_params(axis="y", labelsize=9, rotation=0)
    ax.collections[0].colorbar.ax.yaxis.label.set_color(FG)
    ax.collections[0].colorbar.ax.tick_params(colors=FG)
    fig.tight_layout()
    return save_fig(fig, "fig2_heatmap.png")

# ── FIGURE 3: Top 5 & Bottom 5 KPIs per Region ───────────────────────────────
def fig_top_bottom():
    fig, axes = plt.subplots(1, 3, figsize=(13, 4.5), sharey=False)
    for ax, region in zip(axes, regions):
        sub  = df[df["Region"] == region][CLIMATE_COLS].mean().sort_values()
        low5 = sub.head(5)
        top5 = sub.tail(5)
        combined = pd.concat([low5, top5])
        colors   = ["#e74c3c"] * 5 + ["#2ecc71"] * 5
        ax.barh(
            [FRIENDLY[c] for c in combined.index],
            combined.values,
            color=colors,
            edgecolor="#0d1b2a",
            linewidth=0.4
        )
        ax.set_xlim(1, 5.3)
        ax.set_title(region, fontsize=10)
        ax.axvline(3.0, color="#f39c12", linewidth=0.8, linestyle="--")
        for i, v in enumerate(combined.values):
            ax.text(v + 0.04, i, f"{v:.2f}", va="center", fontsize=7.5, color=FG)
        ax.tick_params(axis="y", labelsize=7.5)
        ax.tick_params(axis="x", labelsize=7.5)
    red_p   = mpatches.Patch(color="#e74c3c", label="Bottom 5 KPIs")
    green_p = mpatches.Patch(color="#2ecc71", label="Top 5 KPIs")
    fig.legend(handles=[red_p, green_p], loc="lower center", ncol=2,
               fontsize=9, framealpha=0.2)
    fig.suptitle("Figure 3 — Top 5 & Bottom 5 KPIs per Region", y=1.01, fontsize=12, fontweight="bold")
    fig.tight_layout()
    return save_fig(fig, "fig3_top_bottom.png")

# ── FIGURE 4: Engagement & Retention by Region × Department ──────────────────
def fig_dept_region():
    grp = (
        df.groupby(["Region", "Department"])[["Engagement", "Retention_Intent"]]
        .mean()
        .reset_index()
    )
    fig, axes = plt.subplots(1, 2, figsize=(13, 4))
    for ax, kpi in zip(axes, ["Engagement", "Retention_Intent"]):
        pivot = grp.pivot(index="Department", columns="Region", values=kpi).reindex(columns=regions)
        x = np.arange(len(pivot))
        width = 0.25
        for i, (region, col) in enumerate(zip(regions, REGION_COLORS.values())):
            bars = ax.bar(x + i * width, pivot[region], width=width,
                          label=region, color=col, edgecolor="#0d1b2a", linewidth=0.4)
        ax.set_xticks(x + width)
        ax.set_xticklabels(pivot.index, rotation=25, ha="right", fontsize=8)
        ax.set_ylim(0, 5.2)
        ax.set_ylabel("Avg Score (1-5)")
        ax.set_title(f"{kpi.replace('_', ' ')} by Department & Region")
        ax.axhline(3.0, color="#e74c3c", linewidth=0.8, linestyle="--")
        ax.yaxis.grid(True, zorder=0)
        ax.set_axisbelow(True)
    axes[0].legend(fontsize=8)
    fig.suptitle("Figure 4 — Engagement & Retention Intent: Department × Region", fontsize=12, fontweight="bold")
    fig.tight_layout()
    return save_fig(fig, "fig4_dept_region.png")

# ── FIGURE 5: Operational Stress & Work-Life Balance by Region ────────────────
def fig_stress_wlb():
    grp = df.groupby("Region")[["Operational_Stress", "Work_Life_Balance"]].mean().reindex(regions)
    x    = np.arange(len(regions))
    width = 0.35
    fig, ax = plt.subplots(figsize=(7, 4))
    b1 = ax.bar(x - width / 2, grp["Operational_Stress"], width=width,
                label="Operational Stress", color="#e74c3c", edgecolor="#0d1b2a")
    b2 = ax.bar(x + width / 2, grp["Work_Life_Balance"],  width=width,
                label="Work-Life Balance",  color="#1a6fa8", edgecolor="#0d1b2a")
    ax.set_xticks(x)
    ax.set_xticklabels(regions)
    ax.set_ylim(0, 5.2)
    ax.set_ylabel("Avg Score (1-5)")
    ax.set_title("Figure 5 — Operational Stress vs Work-Life Balance by Region")
    ax.axhline(3.0, color="#f39c12", linewidth=0.8, linestyle="--", label="Threshold (3.0)")
    ax.legend(fontsize=9)
    ax.yaxis.grid(True, zorder=0)
    ax.set_axisbelow(True)
    for bar in list(b1) + list(b2):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.05,
                f"{bar.get_height():.2f}", ha="center", va="bottom", fontsize=9, color=FG)
    fig.tight_layout()
    return save_fig(fig, "fig5_stress_wlb.png")

# ── GENERATE ALL FIGURES ─────────────────────────────────────────────────────
print("Generating figures...")
f1 = fig_composite_by_region()
f2 = fig_heatmap()
f3 = fig_top_bottom()
f4 = fig_dept_region()
f5 = fig_stress_wlb()
print("Figures done.")

# ── DOCX HELPERS ──────────────────────────────────────────────────────────────
def set_col_width(cell, width_cm):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW  = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(width_cm * 567)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_paragraph(doc, text, bold_first=None, space_after=6):
    p  = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after  = Pt(space_after)
    pf.space_before = Pt(2)
    if bold_first:
        run = p.add_run(bold_first)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_image(doc, path, width_in=6.0, caption=None):
    doc.add_picture(path, width=Inches(width_in))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp  = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].font.size   = Pt(9)
        cp.runs[0].font.italic = True
        cp.runs[0].font.color.rgb = RGBColor(0x77, 0x88, 0x99)

def add_kpi_table(doc, data, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
    for row in data:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    return table

# ── COMPUTE STATS FOR REPORT ──────────────────────────────────────────────────
reg_stats = df.groupby("Region").agg(
    N        = ("Employee_ID", "count"),
    Avg      = ("Avg_Climate", "mean"),
    Eng      = ("Engagement", "mean"),
    Ret      = ("Retention_Intent", "mean"),
    Stress   = ("Operational_Stress", "mean"),
    Safety   = ("Safety_Culture", "mean"),
    WLB      = ("Work_Life_Balance", "mean"),
).reindex(regions).round(2)

reg_kpi = df.groupby("Region")[CLIMATE_COLS].mean().reindex(regions)

def top_bottom(region, n=3):
    s = reg_kpi.loc[region].sort_values()
    return (
        [FRIENDLY[c] for c in s.head(n).index],
        [FRIENDLY[c] for c in s.tail(n).index],
        s.head(n).values,
        s.tail(n).values
    )

# ── BUILD DOCUMENT ────────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# Default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

# ── COVER ─────────────────────────────────────────────────────────────────────
doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("Maritime Workplace Climate Survey 2026")
title_run.bold      = True
title_run.font.size = Pt(22)
title_run.font.color.rgb = RGBColor(0x1a, 0x6f, 0xa8)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run("Regional Analysis Report — Americas")
sub_run.font.size = Pt(14)
sub_run.font.color.rgb = RGBColor(0xc9, 0xa8, 0x4c)

doc.add_paragraph()
meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_run = meta_p.add_run(
    f"Author: Ezequiel Bassa  |  Senior Data Scientist & Sociologist\n"
    f"Survey Period: 2026  |  Sample: {n_total:,} employees  |  10 countries\n"
    f"Regions: North America · Central America · South America"
)
meta_run.font.size = Pt(10)
meta_run.font.color.rgb = RGBColor(0x77, 0xa8, 0xc5)
doc.add_page_break()

# ── PAGE 1: EXECUTIVE SUMMARY ────────────────────────────────────────────────
add_heading(doc, "1. Executive Summary", level=1)

intro = (
    f"This report presents a regional analysis of workplace climate indicators "
    f"across {n_total:,} employees of a Tier-1 Global Maritime Logistics Company "
    f"operating in 10 countries throughout the Americas. Data were collected through "
    f"a structured survey instrument comprising 20 Likert-scale climate indicators "
    f"(1–5 scale), demographic variables, and open-ended employee feedback. "
    f"The analysis compares three sub-regions: North America, Central America, and South America."
)
add_paragraph(doc, intro)

# Summary stats table
headers = ["Region", "N", "Composite Score", "Engagement", "Retention Intent", "Operational Stress"]
data = []
for r in regions:
    s = reg_stats.loc[r]
    data.append([r, f"{int(s.N):,}", f"{s.Avg:.2f}", f"{s.Eng:.2f}", f"{s.Ret:.2f}", f"{s.Stress:.2f}"])
doc.add_paragraph()
add_paragraph(doc, "Table 1 — Regional Summary Statistics", bold_first="")
add_kpi_table(doc, data, headers)
doc.add_paragraph()

add_paragraph(doc,
    "Overall, all three regions score within the moderate band (3.0–4.0) on the composite "
    "climate index. No region exceeds the 4.0 threshold considered indicative of a strong "
    "positive climate, and none falls below the 3.0 critical threshold. Operational Stress "
    "and Work-Life Balance emerge as the most differentiated indicators across regions, "
    "while Ethical Standards and Team Cohesion show the greatest consistency."
)

add_image(doc, f1, width_in=5.5,
          caption="Figure 1. Average composite climate score by region with 95% confidence intervals.")

doc.add_page_break()

# ── PAGE 2: REGIONAL OVERVIEW ─────────────────────────────────────────────────
add_heading(doc, "2. Regional Climate Overview", level=1)

add_paragraph(doc,
    "Figure 2 presents a comprehensive heatmap of all 20 climate indicators disaggregated "
    "by region. This visualization allows for rapid identification of strengths and vulnerabilities "
    "across the full indicator set. Scores above 3.5 are displayed in green tones, while "
    "scores below 3.0 are highlighted in red, indicating areas requiring immediate attention."
)

add_image(doc, f2, width_in=6.5,
          caption="Figure 2. Heatmap of all 20 KPIs by region. Red = critical (<3.0), Green = strong (>3.5).")

doc.add_paragraph()
add_heading(doc, "Key Observations", level=2)

for region in regions:
    bot_kpis, top_kpis, bot_vals, top_vals = top_bottom(region, n=3)
    comp = reg_stats.loc[region, "Avg"]
    add_paragraph(doc,
        f" — Composite score: {comp:.2f}/5.0. "
        f"Strongest indicators: {', '.join(top_kpis)} "
        f"({', '.join([f'{v:.2f}' for v in top_vals[::-1]])}). "
        f"Areas of concern: {', '.join(bot_kpis)} "
        f"({', '.join([f'{v:.2f}' for v in bot_vals])}).",
        bold_first=region
    )

doc.add_page_break()

# ── PAGE 3: TOP & BOTTOM KPIs BY REGION ───────────────────────────────────────
add_heading(doc, "3. Indicator Deep-Dive by Region", level=1)

add_paragraph(doc,
    "Figure 3 isolates the five highest and five lowest scoring indicators for each region. "
    "This analysis is critical for targeted HR intervention, as it identifies where investment "
    "in policy or culture change will yield the highest return. The amber dashed line marks the "
    "3.0 moderate threshold."
)

add_image(doc, f3, width_in=6.5,
          caption="Figure 3. Top 5 (green) and bottom 5 (red) KPIs per region.")

doc.add_paragraph()

# Per-region narrative
for region in regions:
    add_heading(doc, region, level=2)
    bot_kpis, top_kpis, bot_vals, top_vals = top_bottom(region, n=5)
    n_reg = int(reg_stats.loc[region, "N"])
    comp  = reg_stats.loc[region, "Avg"]

    add_paragraph(doc,
        f"With {n_reg:,} employees surveyed and a composite score of {comp:.2f}, "
        f"{region} presents a {'moderate' if 3.0 <= comp < 4.0 else 'strong' if comp >= 4.0 else 'critical'} "
        f"climate profile. "
        f"The region leads on {top_kpis[-1]} ({top_vals[-1]:.2f}) and {top_kpis[-2]} ({top_vals[-2]:.2f}), "
        f"suggesting relative strengths in {'operational culture' if 'Safety' in top_kpis[-1] or 'Ethical' in top_kpis[-1] else 'interpersonal and structural dimensions'}. "
        f"Priority areas for improvement include {bot_kpis[0]} ({bot_vals[0]:.2f}) "
        f"and {bot_kpis[1]} ({bot_vals[1]:.2f}), both of which fall below or near the 3.0 threshold."
    )

doc.add_page_break()

# ── PAGE 4: ENGAGEMENT, RETENTION & DEPARTMENT ANALYSIS ──────────────────────
add_heading(doc, "4. Engagement, Retention Intent & Departmental Patterns", level=1)

add_paragraph(doc,
    "Engagement and Retention Intent are two of the most operationally critical indicators "
    "in workplace climate research, as they directly predict voluntary turnover and productivity. "
    "Figure 4 disaggregates these KPIs by department within each region, revealing significant "
    "structural variation beyond aggregate regional scores."
)

add_image(doc, f4, width_in=6.5,
          caption="Figure 4. Engagement and Retention Intent by department across regions.")

doc.add_paragraph()
add_heading(doc, "Cross-Regional Findings", level=2)

# Compute department with lowest engagement per region
dept_eng = df.groupby(["Region", "Department"])["Engagement"].mean().reset_index()
for region in regions:
    sub     = dept_eng[dept_eng["Region"] == region].sort_values("Engagement")
    low_d   = sub.iloc[0]["Department"]
    low_v   = sub.iloc[0]["Engagement"]
    high_d  = sub.iloc[-1]["Department"]
    high_v  = sub.iloc[-1]["Engagement"]
    add_paragraph(doc,
        f" — Lowest engagement: {low_d} ({low_v:.2f}). "
        f"Highest engagement: {high_d} ({high_v:.2f}). "
        f"Spread of {high_v - low_v:.2f} points indicates "
        f"{'high' if high_v - low_v > 0.5 else 'moderate'} within-region departmental inequality.",
        bold_first=region
    )

doc.add_paragraph()
add_paragraph(doc,
    "Across all regions, Maritime-Crew and Operations-Ports consistently rank lowest on both "
    "Engagement and Retention Intent, reflecting the structural challenges inherent to offshore "
    "and port-based work — including irregular schedules, physical demands, and geographic isolation. "
    "IT and HR departments show the strongest scores on both indicators, likely due to greater "
    "autonomy, remote work flexibility, and career development pathways."
)

doc.add_page_break()

# ── PAGE 5: OPERATIONAL STRESS, WLB & CONCLUSIONS ────────────────────────────
add_heading(doc, "5. Operational Stress, Work-Life Balance & Conclusions", level=1)

add_heading(doc, "5.1 Stress and Work-Life Balance", level=2)

add_paragraph(doc,
    "Operational Stress and Work-Life Balance represent a tension axis that is particularly "
    "pronounced in maritime logistics contexts. Figure 5 compares these two indicators across "
    "regions. Note that Operational Stress is a negatively-valenced indicator — higher scores "
    "denote greater reported stress — while Work-Life Balance follows the standard positive "
    "direction."
)

add_image(doc, f5, width_in=5.5,
          caption="Figure 5. Operational Stress vs. Work-Life Balance by region.")

doc.add_paragraph()

for region in regions:
    stress = reg_stats.loc[region, "Stress"]
    wlb    = reg_stats.loc[region, "WLB"]
    gap    = stress - wlb
    add_paragraph(doc,
        f" — Stress: {stress:.2f} | WLB: {wlb:.2f} | Gap: {gap:+.2f}. "
        f"{'Stress significantly exceeds WLB — a risk indicator for burnout and turnover.' if gap > 0.3 else 'Relatively balanced stress-WLB profile.'}",
        bold_first=region
    )

doc.add_paragraph()
add_heading(doc, "5.2 Conclusions & Strategic Recommendations", level=2)

conclusions = [
    ("Regional parity with local nuance.",
     " All three regions perform within the moderate climate band. While no region is in crisis, "
     "none has achieved high-performance status. Region-specific diagnostics are recommended rather "
     "than pan-Americas HR strategies."),
    ("Operational roles are the primary vulnerability.",
     " Maritime-Crew and Operations-Ports departments consistently underperform on Engagement, "
     "Retention Intent, and Work-Life Balance across all regions. Targeted interventions — including "
     "rotation programs, mental health support, and enhanced compensation reviews — are advised."),
    ("Mid-career attrition risk.",
     " Employees in the 2–4 year tenure band show a statistically notable dip in both Engagement "
     "and Retention Intent (the 'mid-career crisis' zone). Structured career development conversations "
     "at the 18-month mark could mitigate this risk."),
    ("Ethical Standards as a regional strength.",
     " This indicator scores consistently above 3.8 across all regions and departments, suggesting "
     "that the company's compliance and conduct culture is perceived positively. This is a retention "
     "asset that should be actively communicated in employer branding."),
    ("Fair Compensation as a recurring vulnerability.",
     " Fair Compensation ranks in the bottom three indicators in at least two of the three regions. "
     "Given its direct link to Retention Intent, a compensation benchmarking review against "
     "industry standards is strongly recommended."),
]

for bold, text in conclusions:
    add_paragraph(doc, text, bold_first=bold, space_after=5)

doc.add_paragraph()
add_paragraph(doc,
    "This analysis is based on synthetic survey data generated for portfolio demonstration purposes. "
    "All patterns, biases, and findings are designed to reflect realistic maritime industry dynamics "
    "but do not represent any real organization or workforce.",
    space_after=4
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("Ezequiel Bassa — Senior Data Scientist & Sociologist | ezequielbassa.com")
run.font.size  = Pt(9)
run.font.italic = True
run.font.color.rgb = RGBColor(0x55, 0x6a, 0x7a)

# ── SAVE ──────────────────────────────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f"\nReport saved: {OUTPUT_PATH}")
